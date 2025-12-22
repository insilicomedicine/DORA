import io
import logging
import os
import re
from typing import Any, List, Optional, Union

from django.conf import settings
from docx import Document as WordDocument
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section as DocxSection
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from bibliography.models import BibliographyType
from documents.document_export.defs import (
    BIB_ID_CHUNK_ID_PATTERN,
    BIB_ID_CHUNK_ID_REFERENCE_PATTERN,
    COVER_PAGE_DISCLAIMER_STRING,
    PUBMED_URL_PATTERN,
    FontSizes,
    RGBColors,
)
from documents.document_export.defs import DocxPageSizes as sizes
from documents.document_export.utils import (
    format_author_str,
    format_file_name,
    format_inline_reference_for_websearch,
    get_formatted_references,
    get_mermaid_image_stream_and_size,
    is_markdown_heading,
    should_skip_section_title,
)
from documents.models import Document, Section

logger = logging.getLogger(__name__)


class DocxGenerator:
    def __init__(self, output_file: Union[str, io.BytesIO, None], document: Document) -> None:
        self.output_file = output_file
        self.word_doc = WordDocument()
        self._set_styles()
        self.document = document
        self.bib_id_pmid_map = document.bib_id_pmid_map
        self.doc_bibliographies = document.get_bibliographies()
        self.bibliography_mapping = {bib["id"]: bib for bib in self.doc_bibliographies}

    def _get_resource_path(self, resource_name: str) -> str:
        return os.path.join(settings.BASE_DIR, "documents/document_export/resources", resource_name)

    def _set_styles(self) -> None:
        styles = self.word_doc.styles

        styles["Title"].font.size = Pt(FontSizes.TITLE)
        styles["Title"].font.color.rgb = RGBColor(0, 0, 0)
        styles["Title"].font.bold = True
        styles["Title"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        styles["Title"].paragraph_format.space_before = sizes.SPACE_ABOVE_TITLE
        for paragraph_format in styles["Title"].paragraph_format.element.xpath(".//w:pBdr"):
            paragraph_format.getparent().remove(paragraph_format)
        title_style = styles.add_style("DORA Title", WD_STYLE_TYPE.PARAGRAPH)
        title_style.base_style = styles["Title"]
        title_style.font.name = "Arial"
        base_heading_styles = ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]
        for base_style_name in base_heading_styles:
            font_size = FontSizes.get_font_size(base_style_name.replace(" ", "").upper())
            styles[base_style_name].font.size = Pt(font_size)
            styles[base_style_name].font.color.rgb = RGBColor(*RGBColors.BLACK)
            styles[base_style_name].font.bold = True
            # Font can not be set with built-in styles.
            # Instead of "Heading n", use "DORA Heading n" if needed.
            new_style_name = f"DORA {base_style_name}"
            new_style = styles.add_style(new_style_name, WD_STYLE_TYPE.PARAGRAPH)
            new_style.base_style = styles[base_style_name]
            new_style.font.name = "Arial"

        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(FontSizes.NORMAL)

        normal_link_style = styles.add_style("NormalHyperlink", WD_STYLE_TYPE.CHARACTER)
        normal_link_style.base_style = styles["Normal"]
        normal_link_style.font.color.rgb = RGBColor(*RGBColors.DORA_GREEN)

        disclaimer_style = styles.add_style("Disclaimer", WD_STYLE_TYPE.CHARACTER)
        disclaimer_style.font.name = "Arial"
        disclaimer_style.font.size = Pt(FontSizes.FOOTER)
        disclaimer_style.font.italic = True

        disclaimer_link_style = styles.add_style("DisclaimerHyperlink", WD_STYLE_TYPE.CHARACTER)
        disclaimer_link_style.base_style = disclaimer_style
        disclaimer_link_style.font.color.rgb = RGBColor(*RGBColors.BLUE)
        disclaimer_link_style.font.underline = True

        indent_style = styles.add_style("SectionParagraph", WD_STYLE_TYPE.PARAGRAPH)
        indent_style.paragraph_format.first_line_indent = sizes.FIRST_LINE_INDENT
        indent_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        reference_link_style = styles.add_style("ReferenceHyperlink", WD_STYLE_TYPE.CHARACTER)
        reference_link_style.base_style = styles["Normal"]

        reference_domain_style = styles.add_style("ReferenceDomain", WD_STYLE_TYPE.CHARACTER)
        reference_domain_style.base_style = styles["Normal"]
        reference_domain_style.font.color.rgb = RGBColor(*RGBColors.GRAY)

    def _add_hyperlink(self, paragraph: Paragraph, text: str, url: str, style_name: Optional[str]) -> None:
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = Run(OxmlElement("w:r"), paragraph)
        new_run.text = text
        if style_name:
            new_run.style = style_name
        hyperlink.append(new_run._element)
        paragraph._p.append(hyperlink)

    def _replace_text(self, paragraph: Paragraph, new_text: str) -> None:
        new_run = Run(OxmlElement("w:r"), paragraph)
        new_run.text = new_text
        paragraph._p.append(new_run._element)

    def _add_header_image(self, section: DocxSection, header: Any, is_cover_page: bool = True) -> None:
        header.is_linked_to_previous = False
        section.left_margin = sizes.MARGIN_LEFT
        section.right_margin = sizes.MARGIN_LEFT
        image_path = (
            self._get_resource_path("header_image_1.jpg")
            if is_cover_page
            else self._get_resource_path("header_image_2.jpg")
        )
        header_run = header.paragraphs[0].add_run()
        header_run.add_picture(image_path)
        for paragraph in header.paragraphs:
            paragraph.paragraph_format.left_indent = -sizes.MARGIN_LEFT
            paragraph.paragraph_format.space_after = sizes.SPACE_BELOW_HEADER
        section.header_distance = sizes.SPACE_ABOVE_HEADER

    def _add_visual_summary(self) -> None:
        image_stream, image_size = get_mermaid_image_stream_and_size(self.document)
        if not image_stream:
            return

        actual_width, actual_height = image_size
        width_ratio = sizes.MAX_AVAILABLE_WIDTH / actual_width
        height_ratio = sizes.MAX_AVAILABLE_HEIGHT / actual_height
        scale_ratio = min(width_ratio, height_ratio)
        image_width = actual_width * scale_ratio
        image_height = actual_height * scale_ratio

        paragraph = self.word_doc.add_paragraph("")
        paragraph_run = paragraph.add_run()
        paragraph_run.add_picture(image_stream, width=image_width, height=image_height)

        text_paragraph = self.word_doc.add_paragraph("Visual summary")
        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_formatted_paragraph(
        self, contents: Optional[List[dict]], style: Optional[str] = None
    ) -> Paragraph:
        paragraph = self.word_doc.add_paragraph("", style=style)
        if not contents:
            return paragraph
        for content_item in contents:
            if content_item["type"] == "hardBreak":
                paragraph.add_run("\n")
            elif content_item["type"] == "text":
                text = content_item["text"]
                current_run = paragraph.add_run(text)
                if "marks" in content_item:
                    for mark in content_item["marks"]:
                        mark_type = mark["type"]
                        if mark_type == "bold":
                            current_run.bold = True
                        elif mark_type == "italic":
                            current_run.italic = True
                        elif mark_type == "underline":
                            current_run.underline = True
                        elif mark_type == "link":
                            current_run.clear()
                            href = mark["attrs"].get("href")
                            bib_id = mark["attrs"].get("id")
                            link_type = mark["attrs"].get("data-type")
                            custom_text = mark["attrs"].get("data-custom_text")
                            if href[:4] == "http":
                                self._add_hyperlink(paragraph, text, href, "NormalHyperlink")
                            elif href == "#" and bib_id:
                                if pmid := self.bib_id_pmid_map.get(bib_id):
                                    self._add_hyperlink(
                                        paragraph,
                                        text,
                                        PUBMED_URL_PATTERN.format(pmid=pmid),
                                        "NormalHyperlink",
                                    )
                                elif link_type == BibliographyType.WEBSEARCH and custom_text:
                                    self._add_hyperlink(paragraph, text, custom_text, "NormalHyperlink")
            else:
                raise ValueError(f"Unsupported element type: {content_item['type']} in paragraph content")
        return paragraph

    def _add_formatted_list_items(
        self, content: List[dict], is_ordered: bool, is_in_blockquote: bool = False
    ) -> Paragraph:
        for content_item in content:
            if content_item["type"] == "listItem":
                for list_item_content in content_item["content"]:
                    if list_item_content["type"] == "paragraph":
                        paragraph = self._add_formatted_paragraph(list_item_content.get("content"))
                        style = "List Number" if is_ordered else "List Bullet"
                        paragraph.style = style
                        paragraph.style.font.italic = is_in_blockquote
                    elif list_item_content["type"] == "orderedList":
                        self._add_formatted_list_items(
                            list_item_content.get("content"),
                            is_ordered=True,
                            is_in_blockquote=is_in_blockquote,
                        )
                    elif list_item_content["type"] == "bulletList":
                        self._add_formatted_list_items(
                            list_item_content.get("content"),
                            is_ordered=False,
                            is_in_blockquote=is_in_blockquote,
                        )
                    else:
                        raise ValueError(
                            f"Unsupported element type: {list_item_content['type']} in listItem content"
                        )
            else:
                raise ValueError(
                    f"Unsupported element type: {content_item['type']} in bulletList/orderedList content"
                )

    def _add_paragraphs_from_sections(self, sections: List[Section], heading_level: int = 1) -> None:
        heading_level = min(heading_level, 4)
        for section in sections:
            # Skip title only for top-level single section with main_text slug
            if not (heading_level == 1 and should_skip_section_title(sections, section)):
                self.word_doc.add_paragraph(section.title, style=f"Heading {heading_level}")

            if section.result is not None:
                if section.result.get("data_format"):
                    self._process_data_format_content(section.result.get("data_format"))
                elif section.result.get("data"):
                    self._process_raw_data_content(section.result.get("data"))

            if section.get_sub_sections():
                self._add_paragraphs_from_sections(section.get_sub_sections(), heading_level + 1)

    def _process_data_format_content(self, data_format: List[dict]) -> None:
        for item in data_format:
            item_type = item["type"]
            if item_type == "paragraph":
                self._add_formatted_paragraph(item.get("content"), style="SectionParagraph")
            elif item_type == "blockquote":
                self._process_blockquote_content(item["content"])
            elif item_type == "orderedList":
                self._add_formatted_list_items(item["content"], is_ordered=True)
            elif item_type == "bulletList":
                self._add_formatted_list_items(item["content"], is_ordered=False)
            elif item_type == "heading":
                self._process_heading_content(item)
            elif item_type == "table":
                self._process_table_content(item)
            else:
                raise ValueError(f"Unsupported element type: {item_type} in data_format_item")

    def _process_blockquote_content(self, content: List[dict]) -> None:
        for item in content:
            content_type = item["type"]
            if content_type == "paragraph":
                paragraph = self._add_formatted_paragraph(item.get("content"))
                paragraph.style = "Quote"
            elif content_type == "orderedList":
                self._add_formatted_list_items(item["content"], is_ordered=True, is_in_blockquote=True)
            elif content_type == "bulletList":
                self._add_formatted_list_items(item["content"], is_ordered=False, is_in_blockquote=True)
            elif content_type == "heading":
                self._process_heading_content(item)
            else:
                raise ValueError(f"Unsupported element type: {content_type} in blockquote content")

    def _process_heading_content(self, item: dict) -> None:
        text_heading_level = item["attrs"].get("level", 1)
        text_heading_level = min(text_heading_level, 4)
        self._add_formatted_paragraph(item.get("content"), style=f"Heading {text_heading_level}")

    def _process_table_content(self, item: dict) -> None:
        num_cols = self._get_table_column_count(item)
        if num_cols == 0:
            return

        table = self._create_table(num_cols)
        self._fill_table_content(table, item["content"])
        self.word_doc.add_paragraph()  # Add empty line after table

    def _get_table_column_count(self, table_item: dict) -> int:
        num_cols = 0
        for row in table_item.get("content", []):
            if row["type"] == "tableRow":
                num_cols = max(num_cols, len(row.get("content", [])))
        return num_cols

    def _create_table(self, num_cols: int) -> Any:
        table = self.word_doc.add_table(rows=0, cols=num_cols)
        table.style = "Table Grid"
        return table

    def _fill_table_content(self, table: Any, content: List[dict]) -> None:
        for row in content:
            if row["type"] != "tableRow":
                continue

            row_cells = table.add_row().cells
            for idx, cell in enumerate(row.get("content", [])):
                if idx >= len(row_cells):
                    break

                cell_type = cell["type"]
                self._process_cell_content(row_cells[idx], cell)

                # Bold the header cell text
                if cell_type == "tableHeader":
                    for paragraph in row_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    def _process_cell_content(self, cell_obj: Any, cell_data: dict) -> None:
        """Process the cell content including any bibliography links."""
        if cell_data["type"] not in ["tableHeader", "tableCell"]:
            return

        # Clear existing paragraph
        paragraph = cell_obj.paragraphs[0]
        paragraph.clear()

        for para_item in cell_data.get("content", []):
            if para_item["type"] == "paragraph":
                self._process_cell_paragraph_content(paragraph, para_item.get("content", []))

    def _process_cell_paragraph_content(self, paragraph: Paragraph, content: List[dict]) -> None:
        """Process paragraph content inside a table cell, including bibliography links."""
        for content_item in content:
            if content_item["type"] == "hardBreak":
                paragraph.add_run("\n")
            elif content_item["type"] == "text":
                text = content_item["text"]

                # Check for bibliography references
                bib_pattern = re.compile(BIB_ID_CHUNK_ID_REFERENCE_PATTERN)

                if bib_pattern.search(text):
                    # Handle bibliography references
                    self._process_bibliography_references(text, paragraph)
                else:
                    # Normal text processing including marks
                    current_run = paragraph.add_run(text)
                    if "marks" in content_item:
                        for mark in content_item["marks"]:
                            mark_type = mark["type"]
                            if mark_type == "bold":
                                current_run.bold = True
                            elif mark_type == "italic":
                                current_run.italic = True
                            elif mark_type == "underline":
                                current_run.underline = True
                            elif mark_type == "link":
                                current_run.clear()
                                href = mark["attrs"].get("href")
                                bib_id = mark["attrs"].get("id")
                                link_type = mark["attrs"].get("data-type")
                                custom_text = mark["attrs"].get("data-custom_text")
                                if href[:4] == "http":
                                    self._add_hyperlink(paragraph, text, href, "NormalHyperlink")
                                elif href == "#" and bib_id:
                                    if pmid := self.bib_id_pmid_map.get(bib_id):
                                        self._add_hyperlink(
                                            paragraph,
                                            text,
                                            PUBMED_URL_PATTERN.format(pmid=pmid),
                                            "NormalHyperlink",
                                        )
                                    elif link_type == BibliographyType.WEBSEARCH and custom_text:
                                        self._add_hyperlink(paragraph, text, custom_text, "NormalHyperlink")

    def _process_raw_data_content(self, data: str) -> None:
        parts = self._split_content_by_markdown_tables(data)
        for part in parts:
            if self._is_markdown_table(part):
                self._add_markdown_table(part)
            else:
                paragraph_texts = [
                    paragraph_text for paragraph_text in part.split("\n") if paragraph_text.strip()
                ]
                for paragraph_text in paragraph_texts:
                    self._add_paragraphs_with_bib_links(paragraph_text)

    def _get_formatted_paragraph(self, contents: Optional[List[dict]]) -> str:
        """Helper method to format paragraph contents"""
        text = ""
        if not contents:
            return text
        for content_item in contents:
            if content_item["type"] == "text":
                text += content_item["text"]
        return text

    def _add_paragraphs_with_bib_links(self, paragraph_text: str) -> None:
        is_heading, heading_level = is_markdown_heading(paragraph_text)
        if is_heading:
            paragraph_text = paragraph_text.strip("# ")
            paragraph_style = f"Heading {heading_level}"
        else:
            paragraph_style = "SectionParagraph"

        paragraph = self.word_doc.add_paragraph(style=paragraph_style)
        self._process_bibliography_references(paragraph_text, paragraph)

    def _add_text_with_markdown_formatting(self, paragraph: Paragraph, text: str) -> None:
        """
        Add text to paragraph with markdown bold and italic formatting

        Args:
            paragraph: The paragraph object to add text to
            text: The text that may contain markdown bold (**text**) and italic (*text*) syntax
        """
        if not text:
            return

        # Process bold first: **text**
        # Use a placeholder to protect bold text while processing italic
        bold_pattern = re.compile(r"\*\*([^\*]+)\*\*")
        bold_parts = []
        bold_placeholders = {}

        for i, match in enumerate(bold_pattern.finditer(text)):
            placeholder = f"__BOLD_{i}__"
            bold_placeholders[placeholder] = match.group(1)
            bold_parts.append((match.start(), match.end(), placeholder))

        # Replace bold patterns with placeholders (in reverse to maintain indices)
        for start, end, placeholder in reversed(bold_parts):
            text = text[:start] + placeholder + text[end:]

        # Process italic: *text*
        italic_pattern = re.compile(r"\*([^\*]+)\*")
        segments = []
        last_end = 0

        for match in italic_pattern.finditer(text):
            # Add text before italic
            if match.start() > last_end:
                segments.append(("normal", text[last_end : match.start()]))
            # Add italic text
            segments.append(("italic", match.group(1)))
            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            segments.append(("normal", text[last_end:]))

        # If no italic found, treat whole text as normal
        if not segments:
            segments = [("normal", text)]

        # Now process segments and replace bold placeholders
        for seg_type, seg_text in segments:
            # Check for bold placeholders in this segment
            if "__BOLD_" in seg_text:
                # Split by bold placeholders
                parts = re.split(r"(__BOLD_\d+__)", seg_text)
                for part in parts:
                    if part in bold_placeholders:
                        # This is a bold placeholder
                        run = paragraph.add_run(bold_placeholders[part])
                        run.bold = True
                        if seg_type == "italic":
                            run.italic = True
                    elif part:
                        # Regular text
                        run = paragraph.add_run(part)
                        if seg_type == "italic":
                            run.italic = True
            else:
                # No bold placeholders, just add the text
                run = paragraph.add_run(seg_text)
                if seg_type == "italic":
                    run.italic = True

    def _process_bibliography_references(self, text: str, paragraph: Paragraph) -> None:
        """
        Process bibliography references in the text and add them to the specified paragraph

        Args:
            text: The text containing bibliography references
            paragraph: The paragraph object to which the references should be added
        """
        pattern = re.compile(BIB_ID_CHUNK_ID_REFERENCE_PATTERN)
        last_end = 0

        for match in pattern.finditer(text):
            start, end = match.span()
            self._add_text_with_markdown_formatting(paragraph, text[last_end:start])
            ref_text = match.group(0)
            matches = re.findall(BIB_ID_CHUNK_ID_PATTERN, ref_text)
            bib_ids = list(set([match[0] for match in matches]))

            for index, bib_id in enumerate(bib_ids):
                prefix = "; " if index > 0 else ""
                bib_info = self.bibliography_mapping.get(bib_id)
                if not bib_info or not bib_info.get("metadata"):
                    continue

                metadata = bib_info["metadata"]
                if bib_info["type"] == BibliographyType.PUBMED:
                    pmid = metadata.get("pubmed_id")
                    authors_str = format_author_str(metadata)
                    hyperlink = PUBMED_URL_PATTERN.format(pmid=pmid)
                    self._add_hyperlink(paragraph, prefix + authors_str, hyperlink, "NormalHyperlink")
                elif bib_info["type"] == BibliographyType.FILE:
                    file_name = format_file_name(metadata.get("file_name", ""))
                    self._replace_text(paragraph, prefix + file_name)
                elif bib_info["type"] == BibliographyType.WEBSEARCH:
                    authors_str, hyperlink = format_inline_reference_for_websearch(metadata)
                    self._add_hyperlink(paragraph, prefix + authors_str, hyperlink, "NormalHyperlink")

            last_end = end

        # Add remaining text after the last match
        self._add_text_with_markdown_formatting(paragraph, text[last_end:])

    def _split_content_by_markdown_tables(self, content: str) -> List[str]:
        """
        Split content by markdown tables, returning a list of table and non-table parts
        """
        # Ensure content ends with a newline for consistent pattern matching
        if not content.endswith("\n"):
            content = content + "\n"

        # Match table header and separator
        table_start_pattern = re.compile(r"\|[^\n]+\|\n\|[-:| ]+\|")

        parts = []
        last_end = 0

        # Find all potential table starts
        for match in table_start_pattern.finditer(content):
            start = match.start()

            # If there's content before this table, add it
            if start > last_end:
                parts.append(content[last_end:start])

            # Now extract the full table including all rows
            table_text = ""
            lines = content[start:].split("\n")
            table_end_line = 0

            # Add header and separator lines
            table_text += lines[0] + "\n" + lines[1] + "\n"
            table_end_line = 2

            # Add data rows as long as they start with | and contain |
            while (
                table_end_line < len(lines)
                and lines[table_end_line].strip()
                and lines[table_end_line].startswith("|")
                and "|" in lines[table_end_line][1:]
            ):
                table_text += lines[table_end_line] + "\n"
                table_end_line += 1

            # Calculate the position after this table
            last_end = start + len(table_text)
            parts.append(table_text)

        # Add any remaining content
        if last_end < len(content):
            parts.append(content[last_end:])

        return parts

    def _is_markdown_table(self, text: str) -> bool:
        """
        Determine if the text is a markdown table
        """
        lines = text.strip().split("\n")
        if len(lines) < 2:
            return False

        # Check if first line starts and ends with |
        if not (lines[0].startswith("|") and lines[0].endswith("|")):
            return False

        # Check if second line is a separator row
        separator_line = lines[1]
        if not (separator_line.startswith("|") and separator_line.endswith("|")):
            return False

        # Separator row should only contain -:|
        separator_parts = separator_line.strip("|").split("|")
        for part in separator_parts:
            if not re.match(r"^[-:| ]+$", part.strip()):
                return False

        return True

    def _add_markdown_table(self, markdown_table: str) -> None:
        """
        Convert markdown table to Word table
        """
        lines = markdown_table.strip().split("\n")
        if len(lines) < 2:
            return

        # Extract headers
        header_line = lines[0]
        headers = [cell.strip() for cell in header_line.strip("|").split("|")]
        num_cols = len(headers)

        # Extract separator line to determine alignment
        separator_line = lines[1]
        alignments = []
        for sep in separator_line.strip("|").split("|"):
            sep = sep.strip()
            if sep.startswith(":") and sep.endswith(":"):
                alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
            elif sep.endswith(":"):
                alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
            else:
                alignments.append(WD_ALIGN_PARAGRAPH.LEFT)

        # Extract table data rows
        data_rows = []
        for i in range(2, len(lines)):
            if not lines[i].strip():
                continue
            cells = [cell.strip() for cell in lines[i].strip("|").split("|")]
            # Ensure each row has the same number of columns as the header
            while len(cells) < num_cols:
                cells.append("")
            data_rows.append(cells[:num_cols])

        # Create Word table
        table = self.word_doc.add_table(rows=len(data_rows) + 1, cols=num_cols)
        table.style = "Table Grid"

        # Fill headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # Set header style
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignments[i]
                for run in paragraph.runs:
                    run.bold = True

        # Fill data rows
        for i, row_data in enumerate(data_rows):
            row = table.rows[i + 1]  # +1 to skip header row
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]

                # Check for bibliography references first (has higher priority)
                bib_pattern = re.compile(BIB_ID_CHUNK_ID_REFERENCE_PATTERN)
                bib_match = bib_pattern.search(cell_text)

                # Check for markdown links
                link_pattern = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
                link_match = link_pattern.search(cell_text)

                if bib_match:
                    # Handle cell with bibliography reference
                    paragraph = cell.paragraphs[0]
                    paragraph.text = ""  # Clear default text
                    self._process_bibliography_references(cell_text, paragraph)
                elif link_match:
                    # Handle cell with markdown link
                    text, url = link_match.groups()
                    paragraph = cell.paragraphs[0]
                    paragraph.text = ""  # Clear default text
                    self._add_hyperlink(paragraph, text, url, "NormalHyperlink")
                else:
                    # Normal text cell
                    cell.text = cell_text

                # Set cell alignment
                for paragraph in cell.paragraphs:
                    paragraph.alignment = alignments[j if j < len(alignments) else 0]

        # Add empty line after table
        self.word_doc.add_paragraph()

    def generate_cover_page(self) -> None:
        title_str = self.document.title
        template_name_str = self.document.template.name
        formatted_date_str = self.document.updated_at.strftime("%b %d, %Y")
        cover_page_section = self.word_doc.sections[0]
        # TODO: remove these lines and _add_header_image method if header is not needed
        # header = cover_page_section.header
        # self._add_header_image(cover_page_section, header)
        # logo_paragraph = header.add_paragraph()
        # logo_run = logo_paragraph.add_run()
        # logo_run.add_picture(
        #     self._get_resource_path("logo_image.jpg"), width=sizes.WIDTH_LOGO, height=sizes.HEIGHT_LOGO
        # )
        # logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.word_doc.add_paragraph(template_name_str, style="Normal")
        self.word_doc.add_paragraph(title_str, style="DORA Title")
        self.word_doc.add_paragraph(formatted_date_str, style="Normal")

        cover_page_section.different_first_page_header_footer = True
        footer = cover_page_section.first_page_footer
        footer_paragraph = footer.paragraphs[0]
        # Find html <a> tags and convert them to hyperlinks
        hyperlink_pattern = re.compile(r"(<font color='lightblue'><u><a href='.*?'>.*?</a></u></font>)")
        parts = hyperlink_pattern.split(COVER_PAGE_DISCLAIMER_STRING)
        for part in parts:
            if part.startswith("<font color='lightblue'><u><a href="):
                match = re.search(r"<font color='lightblue'><u><a href='(.*?)'>(.*?)</a></u></font>", part)
                if match:
                    url, text = match.groups()
                    self._add_hyperlink(footer_paragraph, text, url, "DisclaimerHyperlink")
            else:
                footer_paragraph.add_run(part, style="Disclaimer")

    def generate_main_pages(self):
        main_section = self.word_doc.add_section(WD_SECTION.NEW_PAGE)
        self._add_visual_summary()
        # self._add_header_image(main_section, main_section.header, is_cover_page=False)
        main_section.footer.is_linked_to_previous = False
        main_section.first_page_footer.is_linked_to_previous = False
        footer = main_section.first_page_footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.clear()
        self._add_paragraphs_from_sections(self.document.get_sections())

    def generate_references_pages(self) -> None:
        references = get_formatted_references(self.doc_bibliographies)
        if not references:
            return
        self.word_doc.add_section(WD_SECTION.NEW_PAGE)
        # reference_section = self.word_doc.add_section(WD_SECTION.NEW_PAGE)
        # self._add_header_image(reference_section, reference_section.header, is_cover_page=False)
        self.word_doc.add_paragraph("References", style="Heading 1")
        for reference in references:
            if reference["type"] == BibliographyType.WEBSEARCH:
                reference_text = reference["title"]
            else:
                reference_text = reference["text"]

            if reference["type"] == BibliographyType.WEBSEARCH:
                paragraph = self.word_doc.add_paragraph()
                self._add_hyperlink(paragraph, reference_text, reference["url"], "ReferenceHyperlink")
                paragraph.add_run(f"\n{reference['domain']}", style="ReferenceDomain")
            else:
                self.word_doc.add_paragraph(reference_text, style="Normal")

    def generate_docx(self) -> None:
        self.generate_cover_page()
        self.generate_main_pages()
        self.generate_references_pages()
        if self.output_file:
            self.word_doc.save(self.output_file)
